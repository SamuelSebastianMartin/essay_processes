#! /usr/bin/env python3

import PyPDF2
import docx
import re


class EssayPdf:
    """When passed the name (or PATH/name) of a PDF file, this class
    provides tools to extract the text and remove all artificial
    linebreaks (\n). Currently, there is no provision for headers,
    title pages or bibliographies.
    """
    def __init__(self, filename):
        self.filename = filename
        self.raw_text = self.extract_pdf_text(filename)
        self.converted_text = self.convert()

    def convert(self):
        self.body_text, self.bibliog = self.split_bib(self.raw_text)
        para_marked_text = self.mark_para_breaks(self.body_text)
        no_line_feed_text = self.strip_new_lines(para_marked_text)
        finessed_text = self.finesse_typography(no_line_feed_text)
        converted_text = self.replace_para_breaks(finessed_text)
        return converted_text

    def extract_pdf_text(self, filename):
        """Extracts the text content from the pdf document that is
        passed to it, returning a string.
        """
        pdf_file_obj = open(self.filename, 'rb')
        pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
        page_accumulator = []
        for n in range(pdf_reader.numPages):
            page_obj = pdf_reader.getPage(n)
            page = (page_obj.extractText())
            page_accumulator.append(page)
        raw_text = '\n'.join(page_accumulator)
        return raw_text

    def split_bib(self, raw_text):
        """If there is a bibliography, it is separated so that it is not
        processed with the main text, and line breaks are preserved.
        """
        searches = []
        expr = re.compile("reference.?\\W?", re.IGNORECASE)
        srch_refs = expr.search(raw_text)
        searches.append(srch_refs)
        expr = re.compile("bibliograph.+?\\W?", re.IGNORECASE)
        srch_bib = expr.search(raw_text)
        searches.append(srch_bib)
        for srch in searches:
            if srch:
                body_end = srch.span()[0] - 1
                body_text = raw_text[0: body_end]
                bib_start = srch.span()[0]
                bibliog = raw_text[bib_start:]
        return body_text, bibliog

    def mark_para_breaks(self, body_text):
        """Attempts to find likely places where the \n character
        represents a genuine new paragraph, or new line.
        Such places are marked so they can be identified later.
        """
        # New line after a full stop (far from secure)
        para_marked_text = body_text.replace(".\n", ".NEW_PARA")
        # Some rules for the default ICC title page:
        if "underline family name" in para_marked_text:
            para_marked_text = para_marked_text.replace(
                    "underline family name)", "underline family name)NEW_PARA"
                    )
        if "Term (1,2 or 3)" in para_marked_text:
            para_marked_text = para_marked_text.replace(
                    "Term (1,2 or 3)", "NEW_PARATerm (1,2 or 3)"
                    )
        if "Course Title" in para_marked_text:
            para_marked_text = para_marked_text.replace(
                    "Course Title", "NEW_PARACourse Title"
                    )
        if "Word Count:" in para_marked_text:
            para_marked_text = para_marked_text.replace(
                    "Word Count:", "NEW_PARAWord Count"
                    )

        text = "Title of assignment question  \
                (you should copy this from the assignment question sheet):"
        if "Title of assignment" in para_marked_text:
            para_marked_text = para_marked_text.replace(
                    "Title of assignment", "NEW_PARATitle of assignment")

        expr = re.compile("lecturer.s name", re.IGNORECASE)
        srch = expr.search(para_marked_text)
        if srch:
            text = srch.group()
            para_marked_text = para_marked_text.replace(
                    text, "NEW_PARALecturer's name"
                    )

        expr = re.compile("word count\\W+?\\d+?[,.]\\d{3}\\W?", re.IGNORECASE)
        srch = expr.search(para_marked_text)
        if srch:
            text = srch.group()
            para_marked_text = para_marked_text.replace(
                    text, text+"NEW_PARANEW_PARA"
                    )

        return para_marked_text

    def strip_new_lines(self, para_marked_text):
        """Replaces all new line characters (\n) with a space."""
        no_line_feed_text = para_marked_text.replace("\n", " ")
        return no_line_feed_text

    def finesse_typography(self, no_line_feed_text):
        """Corrects problems with text encoding, and typical errors
        from conversion. This section can be added to when new
        idiosyncrasies are found.
        """
        doublespace_strip = no_line_feed_text.split()
        single_spaced = ' '.join(doublespace_strip)
        finessed_text = single_spaced.replace("Õ", "'")  # Apostrophes
        finessed_text = finessed_text.replace("Ò", "'")  # Single quote
        finessed_text = finessed_text.replace("Ó", "'")  # Single quote
        finessed_text = finessed_text.replace("Þ", "fi")  # Apple Pages!
        if finessed_text[0] == "!":  # Results from the 'table' format.
            finessed_text = finessed_text[1:]
        return finessed_text

    def replace_para_breaks(self, finessed_text):
        """Replaces all the genuine new paragraph points (found in
        `mark_para_breaks()` above) with a \n character.
        """
        converted_text = finessed_text.replace("NEW_PARA", "\n")
        return converted_text


def write_doc_file(essay_text, bibliog):
    paragraph_list = essay_text.split('\n')
    out_doc = docx.Document()
    for paragraph in paragraph_list:
        p = out_doc.add_paragraph()
        p.add_run(paragraph)
    bibliography = out_doc.add_paragraph()
    bibliography.add_run(bibliog)
    out_doc.save(essay.filename[:-3] + "docx")


def filepicker():
    import tkinter as tk
    from tkinter import filedialog
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename()
    return file_path


input("""
        Select a pdf and it will be converted to a docx file
        with the same name
        Press ENTER to begin""")
file_path = filepicker()
essay = EssayPdf(file_path)
essay.convert()
write_doc_file(essay.converted_text, essay.bibliog)
